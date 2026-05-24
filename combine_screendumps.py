import os
from docx import Document
from docx.shared import Inches

def create_test_report():
    screenshots_dir = "reports/screenshots"
    output_file = "reports/Final_Screendump_Report.docx"

    # Mapping of filenames to simple descriptions
    descriptions = {
        "login_success.png": "TC1 – Valid login with tomsmith / SuperSecretPassword!",
        "survey_success.png": "TC2 – Survey submitted successfully with sample data",
        "login_invalid.png": "TC3 – Invalid login attempt with wronguser / wrongpass",
        "email_invalid.png": "TC4 – Invalid email format entered (abc123)",
        "mobile_invalid.png": "TC5 – Invalid mobile number entered (123ABC)",
        "password_invalid.png": "TC6 – Weak password entered (123)",
        "session_timeout.png": "TC7 – Session timeout after 10s idle, forced logout",
        "sql_injection.png": "TC8 – SQL injection attempt (' OR '1'='1) rejected",
        "xss_prevention.png": "TC9 – XSS attempt (<script>alert('XSS')</script>) rejected",
        "boundary_validation.png": "TC10 – Boundary test with 256‑char username/password rejected"
    }

    doc = Document()
    doc.add_heading("Final Screendump Report", 0)

    items = [(f, descriptions[f]) for f in descriptions if os.path.exists(os.path.join(screenshots_dir, f))]

    # Place two screenshots per page, stacked vertically
    for i in range(0, len(items), 2):
        fname1, desc1 = items[i]
        doc.add_heading(desc1, level=2)
        doc.add_picture(os.path.join(screenshots_dir, fname1), width=Inches(6.5))
        doc.add_paragraph(f"Evidence: {fname1}")

        if i+1 < len(items):
            fname2, desc2 = items[i+1]
            doc.add_heading(desc2, level=2)
            doc.add_picture(os.path.join(screenshots_dir, fname2), width=Inches(6.5))
            doc.add_paragraph(f"Evidence: {fname2}")

        # After two screenshots, insert a page break
        doc.add_page_break()

    doc.save(output_file)
    print(f"Report generated: {output_file}")

if __name__ == "__main__":
    create_test_report()
